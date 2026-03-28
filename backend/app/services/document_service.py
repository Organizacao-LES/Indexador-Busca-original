from __future__ import annotations

import csv
import io
import time
import zipfile
from datetime import date, datetime
from pathlib import Path
from uuid import uuid4

import pdfplumber
from docx import Document as DocxDocument
from fastapi import UploadFile
from sqlalchemy.orm import Session

from app.core.config import settings
from app.domain.document import Document
from app.domain.document_category import DocumentCategory
from app.domain.document_history import DocumentHistory
from app.domain.ingestion_history import IngestionHistory
from app.domain.ingestion_status import IngestionStatus
from app.domain.invalid_document import InvalidDocument
from app.domain.user import User
from app.exceptions.document_exceptions import DocumentNotFoundException, DocumentValidationException
from app.repositories.document_repository import DocumentRepository


class DocumentService:
    def __init__(self, document_repository: DocumentRepository):
        self.document_repository = document_repository
        self.allowed_extensions = {
            extension.lower() for extension in settings.DOCUMENT_ALLOWED_EXTENSIONS
        }
        self.max_size_bytes = settings.DOCUMENT_MAX_FILE_SIZE_MB * 1024 * 1024
        self.storage_dir = Path(settings.DOCUMENT_UPLOAD_DIR)

    def upload_document(
        self,
        db: Session,
        *,
        file: UploadFile,
        category: str,
        uploaded_by: User,
        document_date: date | None = None,
    ) -> dict:
        started_at = time.perf_counter()
        filename = file.filename or ""
        normalized_category = category.strip()

        try:
            extension = self._extract_extension(filename)
            self._validate_extension(extension)
            if not normalized_category:
                raise DocumentValidationException("Informe uma categoria válida para o documento.")

            content = file.file.read()
            self._validate_size(content)
            self._validate_integrity(extension, content)
            extracted_content = self._extract_text(extension, content)
        except DocumentValidationException as exc:
            self._register_invalid_document(db, uploaded_by, filename, exc.detail)
            raise
        finally:
            file.file.close()

        storage_path = self._store_file(content, extension)
        title = Path(filename).stem or f"documento-{uuid4().hex[:8]}"
        category_row = self._get_or_create_category(db, normalized_category)
        status_row = self._get_or_create_status(db, "concluido")

        document = Document(
            cod_categoria=category_row.cod_categoria,
            titulo=title,
            tipo=extension.upper(),
            data_publicacao=self._coerce_document_datetime(document_date),
            ativo=True,
            cod_usuario_criador=uploaded_by.cod_usuario,
        )
        db.add(document)
        db.flush()

        history_document = DocumentHistory(
            cod_documento=document.cod_documento,
            cod_usuario=uploaded_by.cod_usuario,
            numero_versao=1,
            caminho_arquivo=str(storage_path),
            texto_extraido=extracted_content,
            texto_processado=extracted_content,
            versao_ativa=True,
        )
        db.add(history_document)

        ingestion_history = IngestionHistory(
            cod_usuario=uploaded_by.cod_usuario,
            cod_documento=document.cod_documento,
            cod_status_ingestao=status_row.cod_status_ingestao,
            tipo_ingestao="manual",
            mensagem_erro=None,
            tempo_processamento_ms=int((time.perf_counter() - started_at) * 1000),
        )
        db.add(ingestion_history)
        db.commit()

        return self.get_document_payload(db, document.cod_documento)

    def get_document_payload(self, db: Session, document_id: int) -> dict:
        payload = self.document_repository.get_document_payload(db, document_id)
        if payload is None:
            raise DocumentNotFoundException()
        return payload

    def list_ingestion_history(self, db: Session, limit: int = 20) -> list[dict]:
        return self.document_repository.list_ingestion_history(db, limit=limit)

    def list_batch_files(self, db: Session, limit: int = 10) -> list[dict]:
        return self.document_repository.list_batch_files(db, limit=limit)

    def to_upload_response(self, payload: dict) -> dict:
        document_date = payload["document_date"].date() if payload["document_date"] else None
        extracted_characters = len(payload["content"] or "")
        return {
            "id": payload["id"],
            "title": payload["title"],
            "fileName": payload["file_name"],
            "category": payload["category"],
            "type": payload["type"],
            "mimeType": self._guess_mime_type(payload["type"].lower()),
            "sizeBytes": payload["size_bytes"],
            "sizeLabel": self._format_size(payload["size_bytes"]),
            "date": document_date,
            "uploadedAt": payload["uploaded_at"],
            "validated": True,
            "integrityOk": True,
            "hash": "",
            "extracted": True,
            "extractedCharacters": extracted_characters,
        }

    def to_history_response(self, payload: dict) -> dict:
        created_at = payload["uploaded_at"]
        extracted_characters = len(payload["content"] or "")
        details = (
            f"Validado ({payload['type']}) • Extraído {extracted_characters} caracteres • "
            f"Status {payload['ingestion_status']}"
        )
        return {
            "date": created_at.strftime("%Y-%m-%d %H:%M") if created_at else "",
            "file": payload["file_name"],
            "type": "Individual",
            "result": "Sucesso" if payload["ingestion_status"] == "concluido" else "Falha",
            "details": details,
        }

    def to_batch_response(self, payload: dict) -> dict:
        status = "indexed" if payload["ingestion_status"] == "concluido" else "error"
        return {
            "name": payload["file_name"],
            "size": self._format_size(payload["size_bytes"]),
            "status": status,
        }

    def to_details_response(self, payload: dict) -> dict:
        document_date = (
            payload["document_date"].isoformat()
            if payload["document_date"]
            else payload["uploaded_at"].isoformat()
        )
        extracted_characters = len(payload["content"] or "")
        return {
            "id": payload["id"],
            "title": payload["title"],
            "category": payload["category"],
            "type": payload["type"],
            "date": document_date,
            "author": payload["author_name"],
            "format": payload["type"],
            "pages": 1,
            "version": int(payload["version"]),
            "indexedAt": payload["uploaded_at"].isoformat(),
            "size": self._format_size(payload["size_bytes"]),
            "downloadUrl": None,
            "content": payload["content"] or "Pré-visualização indisponível para este formato.",
            "extractedCharacters": extracted_characters,
        }

    def _get_or_create_category(
        self, db: Session, category_name: str
    ) -> DocumentCategory:
        category = (
            db.query(DocumentCategory)
            .filter(DocumentCategory.nome_categoria.ilike(category_name))
            .first()
        )
        if category is not None:
            return category

        category = DocumentCategory(nome_categoria=category_name)
        db.add(category)
        db.flush()
        return category

    def _get_or_create_status(
        self, db: Session, status_name: str
    ) -> IngestionStatus:
        status = (
            db.query(IngestionStatus)
            .filter(IngestionStatus.estado_ingestao == status_name)
            .first()
        )
        if status is not None:
            return status

        status = IngestionStatus(estado_ingestao=status_name)
        db.add(status)
        db.flush()
        return status

    def _register_invalid_document(
        self, db: Session, uploaded_by: User, filename: str, reason: str
    ) -> None:
        invalid_document = InvalidDocument(
            cod_usuario=uploaded_by.cod_usuario,
            nome_arquivo=filename or "arquivo-sem-nome",
            motivo_erro=reason[:255],
        )
        db.add(invalid_document)
        db.commit()

    def _coerce_document_datetime(self, value: date | None) -> datetime | None:
        if value is None:
            return None
        return datetime.combine(value, datetime.min.time())

    def _extract_extension(self, filename: str) -> str:
        extension = Path(filename).suffix.lower().lstrip(".")
        if not extension:
            raise DocumentValidationException("O arquivo enviado não possui extensão suportada.")
        return extension

    def _validate_extension(self, extension: str) -> None:
        if extension not in self.allowed_extensions:
            allowed = ", ".join(sorted(ext.upper() for ext in self.allowed_extensions))
            raise DocumentValidationException(
                f"Tipo de arquivo inválido. Formatos aceitos: {allowed}."
            )

    def _validate_size(self, content: bytes) -> None:
        if not content:
            raise DocumentValidationException("O arquivo enviado está vazio.")
        if len(content) > self.max_size_bytes:
            raise DocumentValidationException(
                f"Arquivo excede o tamanho máximo permitido de {settings.DOCUMENT_MAX_FILE_SIZE_MB} MB."
            )

    def _validate_integrity(self, extension: str, content: bytes) -> None:
        if extension == "pdf":
            if not content.startswith(b"%PDF"):
                raise DocumentValidationException("PDF inválido: cabeçalho não reconhecido.")
            if b"%%EOF" not in content[-2048:]:
                raise DocumentValidationException("PDF inválido: final do arquivo corrompido.")
            return

        if extension == "docx":
            try:
                with zipfile.ZipFile(io.BytesIO(content)) as docx_file:
                    if "word/document.xml" not in docx_file.namelist():
                        raise DocumentValidationException(
                            "DOCX inválido: estrutura do documento não reconhecida."
                        )
            except zipfile.BadZipFile as exc:
                raise DocumentValidationException("DOCX inválido: arquivo corrompido.") from exc
            return

        if extension in {"txt", "csv"}:
            try:
                decoded = content.decode("utf-8")
            except UnicodeDecodeError as exc:
                raise DocumentValidationException(
                    "Arquivo de texto inválido: codificação UTF-8 obrigatória."
                ) from exc

            if extension == "csv":
                try:
                    rows = list(csv.reader(decoded.splitlines()))
                except csv.Error as exc:
                    raise DocumentValidationException(
                        "CSV inválido: estrutura inconsistente."
                    ) from exc
                if not rows:
                    raise DocumentValidationException("CSV inválido: arquivo sem registros.")

    def _extract_text(self, extension: str, content: bytes) -> str:
        if extension == "pdf":
            return self._extract_pdf_text(content)
        if extension == "docx":
            return self._extract_docx_text(content)
        if extension == "csv":
            return self._extract_csv_text(content)
        if extension == "txt":
            return self._extract_plain_text(content)
        raise DocumentValidationException("Não há extrator configurado para esse tipo de arquivo.")

    def _extract_pdf_text(self, content: bytes) -> str:
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = [
                    (page.extract_text() or "").strip()
                    for page in pdf.pages
                ]
        except Exception as exc:
            raise DocumentValidationException("Falha ao extrair texto do PDF.") from exc

        extracted_text = "\n\n".join(page for page in pages if page).strip()
        if not extracted_text:
            raise DocumentValidationException("PDF válido, mas sem texto extraível.")
        return extracted_text[:200000]

    def _extract_docx_text(self, content: bytes) -> str:
        try:
            document = DocxDocument(io.BytesIO(content))
        except Exception as exc:
            raise DocumentValidationException("Falha ao extrair texto do DOCX.") from exc

        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        extracted_text = "\n".join(paragraphs).strip()
        if not extracted_text:
            raise DocumentValidationException("DOCX válido, mas sem texto extraível.")
        return extracted_text[:200000]

    def _extract_plain_text(self, content: bytes) -> str:
        try:
            extracted_text = content.decode("utf-8").strip()
        except UnicodeDecodeError as exc:
            raise DocumentValidationException(
                "Arquivo de texto inválido: codificação UTF-8 obrigatória."
            ) from exc

        if not extracted_text:
            raise DocumentValidationException("TXT válido, mas sem conteúdo textual.")
        return extracted_text[:200000]

    def _extract_csv_text(self, content: bytes) -> str:
        decoded = self._extract_plain_text(content)
        rows = list(csv.reader(decoded.splitlines()))
        flattened_rows = [", ".join(cell.strip() for cell in row if cell.strip()) for row in rows]
        extracted_text = "\n".join(row for row in flattened_rows if row).strip()
        if not extracted_text:
            raise DocumentValidationException("CSV válido, mas sem conteúdo textual.")
        return extracted_text[:200000]

    def _store_file(self, content: bytes, extension: str) -> Path:
        self.storage_dir.mkdir(parents=True, exist_ok=True)
        target = self.storage_dir / f"{uuid4().hex}.{extension}"
        target.write_bytes(content)
        return target

    def _format_size(self, size_bytes: int) -> str:
        if size_bytes < 1024:
            return f"{size_bytes} B"
        if size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        return f"{size_bytes / (1024 * 1024):.1f} MB"

    def _guess_mime_type(self, extension: str) -> str:
        if extension == "pdf":
            return "application/pdf"
        if extension == "docx":
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if extension == "csv":
            return "text/csv"
        return "text/plain"


document_service = DocumentService(DocumentRepository())
